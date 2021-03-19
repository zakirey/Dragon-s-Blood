import re
import matplotlib.pyplot as plt
import matplotlib.pyplot as plt1
import requests
from PIL import Image
import docx
from docx.shared import Cm


def title_author_fchapter(book):
    return_values = []
    author_pattern = re.compile(r"Author: (.*)\n")
    title_pattern = re.compile(r"Title: (.*)\n")
    for line in book.readlines():
        result_author = re.match(author_pattern, line)
        result_title = re.match(title_pattern, line)
        if result_author is not None:
            return_values.append(result_author.string)
        if result_title is not None:
            return_values.append(result_title.string)
    with open("Dragon's blood.txt", "r", encoding="utf-8") as file:
        book1 = file.read().replace('\n', ' ')
        result_fchapter = book1.split('CHAPTER I   A LADY AND A GRIFFIN  ')[1].split("   CHAPTER II")[0]
        return_values.append(result_fchapter)
    return return_values


def paragraphs_of_fchapter(fchapter):
    paragraphs = fchapter.split("  ")
    number_of_words = []
    for paragraph in paragraphs:
        words = paragraph.split(" ")
        if len(words) > 10:
            rounded = round(len(words) / 10) * 10
            number_of_words.append(rounded)
        else:
            rounded = len(words)
            number_of_words.append(rounded)

    sorted_words = sorted(number_of_words)

    counter_dict = {i: sorted_words.count(i) for i in sorted_words}
    print(counter_dict)
    plt1.bar(counter_dict.keys(), counter_dict.values())
    plt1.xlabel("Number of words")
    plt1.ylabel("Number of repetitions")
    plt1.savefig("book_counter.jpg")

    plt.plot(sorted_words)
    plt.ylabel("Number of words")
    plt.xlabel("Number of Paragraphs")
    plt.savefig("book_info.jpg")

    return sorted_words


def download_book_cover():
    response = requests.get("https://images-na.ssl-images-amazon.com/images/I/61ht1EMhwmL.jpg")
    file = open("book_cover.jpg", "wb")
    file.write(response.content)
    file.close()


def crop_resize():
    image = Image.open("book_cover.jpg")
    left = 100
    upper = 70
    right = 800
    lower = 1200
    image = image.crop((left, upper, right, lower))
    image = image.resize((400, 600))
    image.save("book_cover.jpg")


def paste_image():
    image = Image.open("logo.gif")
    image1 = Image.open("book_cover.jpg")
    image1_copy = image1.copy()
    image2 = image.rotate(-30, expand=True)
    image2 = image2.resize((50, 50))
    image2_copy = image2.copy()
    image1_copy.paste(image2_copy, (0, 0))
    image1_copy.save("updated_book_cover.jpg")


def create_docx(title, author, sorted_words):
    doc = docx.Document()
    doc.add_paragraph(title, "Title")
    doc.add_picture("updated_book_cover.jpg", width=Cm(11))
    doc.add_heading(author, 1)
    doc.add_paragraph("Report by Zakir Hasanli", "Subtitle")
    doc.add_heading("Chart of distribution of lengths of paragraphs", 1)
    doc.add_picture("book_info.jpg", width=Cm(15))
    doc.add_paragraph("The Plot illustrates how many words were used in each paragraph of the first chapter from "
                      "Dragon's blood book")
    doc.add_paragraph("Number of Paragraphs in the first chapter is " + str(len(sorted_words)), "Caption")
    doc.add_paragraph("Number of words in the first chapter is equal to " + str(sum(sorted_words)), "Caption")
    doc.add_paragraph("Biggest amount of words in one paragraph from first chapter: " + str(sorted_words[-1]),
                      "Caption")
    doc.add_paragraph("Lowest amount of words in one paragraph from first chapter: " + str(sorted_words[0]),
                      "Caption")
    doc.add_paragraph("Average amount of words in paragraph from first chapter: " + str(sum(sorted_words) /
                                                                                        len(sorted_words)) +
                      "\n\n\n\n\n\n\n\n", "Caption")
    doc.add_picture("book_counter.jpg", width=Cm(15))
    doc.add_paragraph("The book was taken from Project Gutenberg: https://www.gutenberg.org/ebooks/10321", "Quote")
    doc.save('book_report.docx')



def run():
    book = open("Dragon's blood.txt", "r", encoding="utf-8")
    a, b, c = title_author_fchapter(book)
    d = paragraphs_of_fchapter(c)
    download_book_cover()
    crop_resize()
    paste_image()
    create_docx(a, b, d)


if __name__ == "__main__":
    run()
